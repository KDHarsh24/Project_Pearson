import io
import os
import hashlib
from typing import List, Dict
from pypdf import PdfReader
from fastapi import UploadFile
from vectorstores.chroma_store import ChromaVectorStore
from db.database import SessionLocal
from db.models import Document, Chunk
import json

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
    """Simple text chunking utility"""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - overlap
    
    return chunks

CASE_FILES_COLLECTION = os.getenv("CASE_FILES_COLLECTION", "case_files")
case_files_store = ChromaVectorStore(collection_name=CASE_FILES_COLLECTION)

SUPPORTED_TEXT_TYPES = {"text/plain"}
SUPPORTED_PDF_TYPES = {"application/pdf"}
SUPPORTED_DOCX_TYPES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}

try:
    import docx  # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

MAX_DOC_LENGTH = 50_000
RAW_DIR = os.getenv("RAW_STORAGE", "storage/raw")
CURATED_DIR = os.getenv("CURATED_STORAGE", "storage/curated")
os.makedirs(RAW_DIR, exist_ok=True)
os.makedirs(CURATED_DIR, exist_ok=True)


def _hash_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:16]


def _extract_text_from_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages_text)


def _extract_text_from_docx(content: bytes) -> str:
    if not _DOCX_AVAILABLE:
        return ""
    bio = io.BytesIO(content)
    document = docx.Document(bio)
    return "\n".join(p.text for p in document.paragraphs)


def ingest_upload(file: UploadFile) -> Dict:
    raw = file.file.read()
    file.file.seek(0)
    file_hash = _hash_bytes(raw)
    filename = file.filename or f"uploaded_{file_hash}.dat"
    content_type = file.content_type or "application/octet-stream"
    if content_type in SUPPORTED_PDF_TYPES or filename.lower().endswith(".pdf"):
        text = _extract_text_from_pdf(raw)
        source_type = "upload_pdf"
    elif (content_type in SUPPORTED_DOCX_TYPES or filename.lower().endswith(".docx")) and _DOCX_AVAILABLE:
        text = _extract_text_from_docx(raw)
        source_type = "upload_docx"
    elif content_type in SUPPORTED_TEXT_TYPES:
        text = raw.decode("utf-8", errors="ignore")
        source_type = "upload_text"
    else:
        return {"status": "ignored", "reason": f"Unsupported type {content_type}"}
    text = text[:MAX_DOC_LENGTH]
    raw_path = os.path.join(RAW_DIR, f"{file_hash}_{filename}")
    if not os.path.exists(raw_path):
        with open(raw_path, "wb") as f:
            f.write(raw)
    
    # Simple document metadata without enrichment
    doc_meta_enriched = {"filename": filename, "content_type": content_type}
    chunks = chunk_text(text, chunk_size=3000, overlap=250)
    metadatas = []
    chunk_meta_jsons = []
    for i, ch in enumerate(chunks):
        # Simple chunk metadata without enrichment
        per_chunk = {"chunk_text": ch[:100] + "...", "chunk_length": len(ch)}
        base_meta = {"filename": filename, "hash": file_hash, "chunk_index": i, "source_type": source_type}
        base_meta.update(per_chunk)
        metadatas.append(base_meta)
        chunk_meta_jsons.append(json.dumps(per_chunk))
    ids = case_files_store.add_texts(chunks, metadatas=metadatas)
    db = SessionLocal()
    try:
        meta_json = {"size": len(raw)}
        meta_json.update(doc_meta_enriched)
        doc = Document(
            filename=filename,
            hash=file_hash,
            content_type=content_type,
            source=source_type,
            raw_path=raw_path,
            meta_json=json.dumps(meta_json)
        )
        db.add(doc)
        db.flush()
        chunk_rows = [
            Chunk(document_id=doc.id, chunk_index=i, text=chunk_text_value, vector_id=vector_id, meta_json=chunk_meta_jsons[i])
            for i, (chunk_text_value, vector_id) in enumerate(zip(chunks, ids))
        ]
        db.add_all(chunk_rows)
        db.commit()
    except Exception as e:
        db.rollback()
        return {"status": "error", "error": str(e)}
    finally:
        db.close()
    return {"status": "ok", "filename": filename, "chunks": len(chunks), "vector_ids": ids, "hash": file_hash, "enriched": True}


def ingest_crawled_case(case_id: str, text: str, source_url: str) -> Dict:
    raw_bytes = text.encode("utf-8", errors="ignore")
    file_hash = _hash_bytes(raw_bytes)
    filename = f"case_{case_id}.txt"
    text = text[:MAX_DOC_LENGTH]
    raw_path = os.path.join(RAW_DIR, f"{file_hash}_{filename}")
    if not os.path.exists(raw_path):
        with open(raw_path, "wb") as f:
            f.write(raw_bytes)
    
    # Simple document metadata without enrichment
    doc_meta_enriched = {"filename": filename, "case_id": case_id, "source_url": source_url}
    chunks = chunk_text(text, chunk_size=3000, overlap=250)
    metadatas = []
    chunk_meta_jsons = []
    for i, ch in enumerate(chunks):
        # Simple chunk metadata without enrichment
        per_chunk = {"chunk_text": ch[:100] + "...", "chunk_length": len(ch)}
        base_meta = {"filename": filename, "hash": file_hash, "chunk_index": i, "source_type": "crawl_case_law", "url": source_url}
        base_meta.update(per_chunk)
        metadatas.append(base_meta)
        chunk_meta_jsons.append(json.dumps(per_chunk))
    ids = case_files_store.add_texts(chunks, metadatas=metadatas)
    db = SessionLocal()
    try:
        meta_json = {"size": len(raw_bytes), "url": source_url}
        meta_json.update(doc_meta_enriched)
        doc = Document(
            filename=filename,
            hash=file_hash,
            content_type="text/plain",
            source="crawl",
            raw_path=raw_path,
            meta_json=json.dumps(meta_json)
        )
        db.add(doc)
        db.flush()
        chunk_rows = [
            Chunk(document_id=doc.id, chunk_index=i, text=chunk_text_value, vector_id=vector_id, meta_json=chunk_meta_jsons[i])
            for i, (chunk_text_value, vector_id) in enumerate(zip(chunks, ids))
        ]
        db.add_all(chunk_rows)
        db.commit()
    except Exception as e:
        db.rollback()
        return {"status": "error", "error": str(e)}
    finally:
        db.close()
    return {"status": "ok", "filename": filename, "chunks": len(chunks), "vector_ids": ids, "hash": file_hash, "enriched": True, "url": source_url}
